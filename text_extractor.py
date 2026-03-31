"""
Text Extraction Module
Handles PDF, DOCX, and plain text file extraction.
"""

import io
import re
import string
import nltk
import pdfplumber
from docx import Document

# Download NLTK data if not present
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

try:
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('wordnet', quiet=True)

from nltk.corpus import stopwords


def extract_text_from_pdf(file) -> str:
    """Extract text from a PDF file object."""
    text = ""
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")
    return text.strip()


def extract_text_from_docx(file) -> str:
    """Extract text from a DOCX file object."""
    text = ""
    try:
        doc = Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")
    return text.strip()


def extract_text_from_txt(file) -> str:
    """Extract text from a plain text file."""
    try:
        content = file.read()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="ignore")
        return content.strip()
    except Exception as e:
        raise ValueError(f"Error reading TXT: {str(e)}")


def extract_text(file, filename: str) -> str:
    """
    Auto-detect file type and extract text.
    Supports: PDF, DOCX, DOC, TXT
    """
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file)
    elif filename_lower.endswith(".txt"):
        return extract_text_from_txt(file)
    else:
        # Try plain text as fallback
        return extract_text_from_txt(file)


def preprocess_text(text: str) -> str:
    """
    Preprocess raw text:
    - Lowercase
    - Remove special characters
    - Remove extra whitespace
    """
    text = text.lower()
    # Keep alphanumeric, spaces, hyphens, dots (for abbreviations like C++)
    text = re.sub(r'[^\w\s\+\#\.\-\/]', ' ', text)
    # Collapse whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def get_stopwords() -> set:
    """Return English stopwords set."""
    try:
        return set(stopwords.words('english'))
    except Exception:
        return {
            'a', 'an', 'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to',
            'for', 'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were',
            'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did',
            'will', 'would', 'could', 'should', 'may', 'might', 'shall', 'can',
            'not', 'no', 'nor', 'so', 'yet', 'both', 'either', 'neither',
            'i', 'we', 'you', 'he', 'she', 'it', 'they', 'this', 'that'
        }
